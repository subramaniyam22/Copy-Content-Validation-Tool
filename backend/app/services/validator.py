import json
import re
import os
from typing import List, Dict, Any, Optional
from fastapi import UploadFile
import pdfplumber
from docx import Document
from openai import AzureOpenAI
from ..config import settings

class ValidatorService:
    def __init__(self):
        self.client = AzureOpenAI(
            api_key=settings.OPENAI_API_KEY,
            api_version=settings.API_VERSION,
            azure_endpoint=settings.AZURE_ENDPOINT
        )

    def parse_guidelines(self, files: List[UploadFile]) -> str:
        """Parses uploaded guideline files (PDF/DOCX) into a single text string."""
        combined_text = ""
        for file in files:
            ext = os.path.splitext(file.filename)[1].lower()
            try:
                if ext == '.pdf':
                    with pdfplumber.open(file.file) as pdf:
                        text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                        combined_text += f"\n--- Limit: {file.filename} ---\n{text}"
                elif ext == '.docx':
                    doc = Document(file.file)
                    text = '\n'.join(para.text for para in doc.paragraphs)
                    combined_text += f"\n--- Limit: {file.filename} ---\n{text}"
                elif ext == '.txt':
                     combined_text += f"\n--- Limit: {file.filename} ---\n{file.file.read().decode('utf-8')}"
            except Exception as e:
                print(f"Error parsing {file.filename}: {e}")
                # Continue processing other files
        return combined_text

    async def validate_grammar(self, scraped_content: List[Dict], guidelines_text: str = None) -> List[Dict]:
        """
        Validates scraped content against guidelines using Azure OpenAI.
        If guidelines_text is provided, it uses that instead of the default JSON guide.
        """
        # Load default prompt (could be moved to DB or config)
        # For now, we'll keep the prompt structure but inject the dynamic guide
        
        system_prompt = "You are a content validator. Analyze the content for spelling, grammar, and adherence to the provided style guide. Provide the output in JSON format."
        if guidelines_text:
             system_prompt += f"\n\nSTYLE GUIDE:\n{guidelines_text}"
        
        # We might want to load the default prompt structure if no custom guide is usually better
        # For this refactor, let's assume if guidelines are uploaded, we use them.
        # If not, we might need a fallback or the user MUST upload them.
        # Given the requirement "upload... guidelines", we prioritize that.
        
        results = []
        for page in scraped_content:
            try:
                paragraphs = page.get("paragraphs", [])
                content_text = "\n".join(p.get("content", "") for p in paragraphs)
                
                if not content_text.strip():
                    continue

                user_message = json.dumps({
                    "page_name": page.get("page_name"),
                    "url": page.get("url"),
                    "content": content_text
                })

                response = self.client.chat.completions.create(
                    model=settings.MODEL_NAME,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_message}
                    ],
                    temperature=0.0,
                    response_format={"type": "json_object"} 
                )
                
                response_content = response.choices[0].message.content
                analysis = json.loads(response_content)
                
                # normalize output structure
                if "schema" in analysis: analysis = analysis["schema"]
                
                # Ensure page_name and url are present (LLM might omit them)
                analysis["page_name"] = page.get("page_name")
                analysis["url"] = page.get("url")
                
                results.append(analysis)

            except Exception as e:
                print(f"Validation error for {page.get('page_name')}: {e}")
                results.append({
                    "page_name": page.get("page_name"),
                    "url": page.get("url"),
                    "errors": [{"type": "system_error", "message": str(e)}]
                })
        
        return results

    async def validate_lv(self, scraped_content: List[Dict], copy_text: str) -> List[Dict]:
        """
        Validates Liquid Variables by comparing Copy Text (source truth) with Scraped Content.
        """
        # This requires the complex logic from validate_lv.py
        # For brevity in this refactor step, I will simplify or copy the logic 
        # structure but mapped to the new class methods.
        
        # 1. Parse Copy Text to JSON (Headings/Content/LVs) using AI
        copy_json = await self._convert_copy_to_json(copy_text)
        
        # 2. Match Content
        validation_results = await self._match_content(copy_json, scraped_content)
        
        return validation_results

    async def _convert_copy_to_json(self, copy_text: str) -> List[Dict]:
        """Convert copy text to JSON format using OpenAI API"""
        prompt = f"""
        You are a content structuring expert. Convert the following copy text into a JSON format with the structure:
        [
          {{
            "page_name": "Page Name",
            "paragraphs": [
              {{
                "heading": "Heading Text",
                "content": "Paragraph content"
              }}
            ]
          }}
        ]

        Rules:
        - Page names are: HOME, FLOOR PLANS, PHOTO GALLERY, AMENITIES, NEIGHBORHOOD (case-sensitive).
        - Headings are marked with **text**, in all caps, or match the pattern 'Welcome to {{location_name}}'.
        - Headings should not be labeled as H1 or H2 in the text (e.g., 'H1:', 'H2:').
        - Ignore lines starting with 'CTA:' and their associated content completely; they do not count as headings or content.
        - A page ends when a new page name is encountered or the text ends.
        - Ignore all content between '[flip cards]' and the next line labeled 'H1:' or 'H2:', including any text in that section.
        - Content is all text following a heading until the next heading, page name, or ignored section (e.g., [flip cards], CTA), concatenated into a single string.
        - Skip lines matching '[.*]' (except page names), '{{button}}', '[lead form]', '[photo cards]', or '[social media & contact info]'.
        - If text follows 'H1:' or 'H2:', treat it as content under the previous heading unless it matches the heading criteria (bold, caps, or 'Welcome to {{location_name}}').
        - Include all paragraphs under each heading, not just the first.
        - Do not include a heading in the JSON if there is no valid content below it (i.e., no text before the next heading, page name, no content/empty content, or ignored section).

        Copy text:
        {copy_text[:15000]} 

        Return the JSON as a string, wrapped in square brackets, e.g., [{...}].
        """
        try:
             response = self.client.chat.completions.create(
                model=settings.MODEL_NAME, 
                messages=[{"role": "system", "content": "You are a content structuring expert that converts text to JSON."}, {"role": "user", "content": prompt}],
                temperature=0.1,
                 response_format={"type": "json_object"}
            )
             
             content = response.choices[0].message.content
             # Handle case where model might return wrapped markdown code block despite json_object mode
             if "```json" in content:
                 content = content.replace("```json", "").replace("```", "")
                 
             match = re.search(r'\[.*\]', content, re.DOTALL)
             if match:
                 return json.loads(match.group())
             return json.loads(content) # Fallback to direct load
             
        except Exception as e:
            print(f"Error parsing copy text: {e}")
            return []

    async def _match_content(self, copy_json: List[Dict], scraped_content: List[Dict]) -> List[Dict]:
        """Improved version that handles adjacent liquid variables better."""
        results = []
        
        # Iterate through each page in the copy data
        for copy_page in copy_json:
            page_name = copy_page['page_name'].upper()
            copy_paragraphs = copy_page.get('paragraphs', [])
            
            print(f"Processing page: {page_name}")
            
            # Find the corresponding page in the scraped data
            scraped_page = next((page for page in scraped_content if page['page_name'].upper() == page_name), None)
            
            if not scraped_page:
                print(f"Page {page_name} not found in scraped data")
                # Handle missing pages
                for copy_para in copy_paragraphs:
                    heading_contexts = self._extract_liquid_variable_contexts(copy_para.get('heading', ''))
                    content_contexts = self._extract_liquid_variable_contexts(copy_para.get('content', ''))
                    
                    for context in heading_contexts + content_contexts:
                        results.append({
                            'pagename': page_name,
                            'context': context['full_context'],
                            'lv': context['variable'],
                            'changed_text': 'NOT_FOUND',
                            'boolean': False
                        })
                continue
            
            scraped_paragraphs = scraped_page.get('paragraphs', [])
            
            # Process each paragraph
            for copy_para in copy_paragraphs:
                copy_heading = copy_para.get('heading', '')
                copy_content = copy_para.get('content', '')
                
                # Process heading liquid variables
                if copy_heading:
                    heading_contexts = self._extract_liquid_variable_contexts(copy_heading)
                    
                    for context in heading_contexts:
                        matched = False
                        
                        # First try: Use multiple variable processing if there are multiple variables
                        if len(heading_contexts) > 1:
                            for scraped_para in scraped_paragraphs:
                                scraped_heading = scraped_para.get('heading', '')
                                if not scraped_heading:
                                    continue
                                
                                multi_results = self._process_multiple_liquid_variables(copy_heading, scraped_heading)
                                for result in multi_results:
                                    if result['variable'] == context['variable'] and result['success']:
                                        results.append({
                                            'pagename': page_name,
                                            'context': context['full_context'],
                                            'lv': context['variable'],
                                            'changed_text': result['replacement'],
                                            'boolean': True
                                        })
                                        matched = True
                                        break
                                if matched:
                                    break
                        
                        # Second try: Individual variable processing
                        if not matched:
                            for scraped_para in scraped_paragraphs:
                                scraped_heading = scraped_para.get('heading', '')
                                if not scraped_heading:
                                    continue
                                
                                # Use fuzzy matching for better results
                                before_match = not context['before'] or self._fuzzy_text_match(context['before'], scraped_heading)
                                after_match = not context['after'] or self._fuzzy_text_match(context['after'], scraped_heading)
                                
                                if before_match and after_match:
                                    # Try to extract the replaced text
                                    replaced_text = self._extract_replacement_text(
                                        scraped_heading, context['before'], context['after']
                                    )
                                    
                                    if replaced_text:
                                        results.append({
                                            'pagename': page_name,
                                            'context': context['full_context'],
                                            'lv': context['variable'],
                                            'changed_text': replaced_text,
                                            'boolean': True
                                        })
                                        matched = True
                                        break
                        
                        if not matched:
                            results.append({
                                'pagename': page_name,
                                'context': context['full_context'],
                                'lv': context['variable'],
                                'changed_text': 'HEADING_MISMATCH',
                                'boolean': False
                            })
                
                # Process content liquid variables
                if copy_content:
                    content_contexts = self._extract_liquid_variable_contexts(copy_content)
                    
                    for context in content_contexts:
                        matched = False
                        
                        # First try: Use multiple variable processing if there are multiple variables
                        if len(content_contexts) > 1:
                            for scraped_para in scraped_paragraphs:
                                scraped_content = scraped_para.get('content', '')
                                if not scraped_content:
                                    continue
                                
                                multi_results = self._process_multiple_liquid_variables(copy_content, scraped_content)
                                for result in multi_results:
                                    if result['variable'] == context['variable'] and result['success']:
                                        results.append({
                                            'pagename': page_name,
                                            'context': context['full_context'],
                                            'lv': context['variable'],
                                            'changed_text': result['replacement'],
                                            'boolean': True
                                        })
                                        matched = True
                                        break
                                if matched:
                                    break
                        
                        # Second try: Individual variable processing
                        if not matched:
                            for scraped_para in scraped_paragraphs:
                                scraped_content = scraped_para.get('content', '')
                                if not scraped_content:
                                    continue
                                
                                # Use fuzzy matching for better results
                                before_match = not context['before'] or self._fuzzy_text_match(context['before'], scraped_content)
                                after_match = not context['after'] or self._fuzzy_text_match(context['after'], scraped_content)
                                
                                if before_match and after_match:
                                    # Try to extract the replaced text
                                    replaced_text = self._extract_replacement_text(
                                        scraped_content, context['before'], context['after']
                                    )
                                    
                                    if replaced_text:
                                        results.append({
                                            'pagename': page_name,
                                            'context': context['full_context'],
                                            'lv': context['variable'],
                                            'changed_text': replaced_text,
                                            'boolean': True
                                        })
                                        matched = True
                                        break
                        
                        if not matched:
                            results.append({
                                'pagename': page_name,
                                'context': context['full_context'],
                                'lv': context['variable'],
                                'changed_text': 'CONTENT_MISMATCH',
                                'boolean': False
                            })
        
        # Validate structure
        validated_result = []
        for item in results:
            if all(key in item for key in ["pagename", "context", "lv", "changed_text", "boolean"]):
                validated_result.append(item)
        
        return validated_result

    def extract_text_from_file(self, file: UploadFile) -> str:
        """Helper to extract text from a single UploadFile"""
        ext = os.path.splitext(file.filename)[1].lower()
        content = ""
        try:
            if ext == '.pdf':
                with pdfplumber.open(file.file) as pdf:
                    content = '\n'.join(page.extract_text() or '' for page in pdf.pages)
            elif ext == '.docx':
                doc = Document(file.file)
                content = '\n'.join(para.text for para in doc.paragraphs)
            elif ext == '.txt':
                content = file.file.read().decode('utf-8')
        except Exception as e:
             print(f"Error reading file {file.filename}: {e}")
        return content

    def _extract_liquid_variable_contexts(self, text):
        """Extract liquid variables with their surrounding context."""
        contexts = []
        if not text or not isinstance(text, str):
            return contexts
        
        pattern = r'{{(\w+)}}'
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        
        if not matches:
            return contexts
        
        for i, match in enumerate(matches):
            var_name = match.group(1)
            start_pos = match.start()
            end_pos = match.end()
            
            context_start = max(0, start_pos - 100)
            sentence_start_match = re.search(r'[.!?]\s+', text[context_start:start_pos])
            if sentence_start_match:
                context_start = context_start + sentence_start_match.end()
            
            context_end = min(len(text), end_pos + 100)
            sentence_end_match = re.search(r'[.!?]', text[end_pos:context_end])
            if sentence_end_match:
                context_end = end_pos + sentence_end_match.start() + 1
            
            full_context = text[context_start:context_end].strip()
            before_text = text[context_start:start_pos]
            after_text = text[end_pos:context_end]
            
            before_clean = re.sub(r'{{.*?}}', '', before_text).strip()
            after_clean = re.sub(r'{{.*?}}', '', after_text).strip()
            
            contexts.append({
                'variable': f'{{{{{var_name}}}}}',
                'full_context': full_context,
                'before': before_clean,
                'after': after_clean,
                'position': (start_pos, end_pos)
            })
        return contexts

    def _fuzzy_text_match(self, needle, haystack, threshold=0.7):
        """Fuzzy text matching to handle minor variations."""
        if not needle or not haystack:
            return True
        
        needle_clean = re.sub(r'\s+', ' ', needle.lower().strip())
        haystack_clean = re.sub(r'\s+', ' ', haystack.lower().strip())
        
        if needle_clean in haystack_clean:
            return True
        
        needle_words = needle_clean.split()
        haystack_words = haystack_clean.split()
        
        if len(needle_words) == 0:
            return True
        
        matches = sum(1 for word in needle_words if word in haystack_words)
        return (matches / len(needle_words)) >= threshold

    def _extract_replacement_text(self, text, before, after):
        """Extract replacement text between before and after markers."""
        try:
            text_clean = re.sub(r'\s+', ' ', text.strip())
            before_clean = re.sub(r'\s+', ' ', before.strip()) if before else ''
            after_clean = re.sub(r'\s+', ' ', after.strip()) if after else ''
            
            if not before_clean and not after_clean:
                return text_clean
            
            if before_clean and after_clean:
                before_pattern = re.escape(before_clean)
                after_pattern = re.escape(after_clean)
                
                pattern = f"{before_pattern}\\s+(.+?)\\s+{after_pattern}"
                match = re.search(pattern, text_clean, re.IGNORECASE)
                if match:
                    return match.group(1).strip()
                    
                pattern = f"{before_pattern}[\\s,]*(.+?)[\\s,]*{after_pattern}"
                match = re.search(pattern, text_clean, re.IGNORECASE)
                if match:
                    return match.group(1).strip()
                    
            elif before_clean:
                before_pattern = re.escape(before_clean)
                pattern = f"{before_pattern}[\\s,]*(.+?)$"
                match = re.search(pattern, text_clean, re.IGNORECASE)
                if match:
                    return match.group(1).strip()
                    
            elif after_clean:
                after_pattern = re.escape(after_clean)
                pattern = f"^(.+?)[\\s,]*{after_pattern}"
                match = re.search(pattern, text_clean, re.IGNORECASE)
                if match:
                    return match.group(1).strip()
            
            return None
        except Exception as e:
            print(f"Error extracting replacement text: {str(e)}")
            return None

    def _process_multiple_liquid_variables(self, copy_text, scraped_text):
        """Process multiple liquid variables in the same text segment."""
        lv_pattern = r'{{(\w+)}}'
        variables = re.findall(lv_pattern, copy_text, re.IGNORECASE)
        
        if not variables:
            return []
        
        results = []
        pattern = copy_text
        for i, var in enumerate(variables):
            pattern = re.sub(f'{{{{{re.escape(var)}}}}}', f'([^.!?]+?)', pattern, flags=re.IGNORECASE)
        
        match = re.search(pattern, scraped_text, re.IGNORECASE | re.DOTALL)
        if match and len(match.groups()) == len(variables):
            for i, var in enumerate(variables):
                replacement = match.group(i + 1).strip()
                replacement = re.sub(r'\s+', ' ', replacement).strip()
                results.append({
                    'variable': f'{{{{{var}}}}}',
                    'replacement': replacement,
                    'success': bool(replacement)
                })
        else:
            for var in variables:
                results.append({
                    'variable': f'{{{{{var}}}}}',
                    'replacement': 'CONTENT_MISMATCH',
                    'success': False
                })
        
        return results

